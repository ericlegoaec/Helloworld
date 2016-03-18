import fl_util
import csv
import datetime
from docx import Document
from xlrd import open_workbook
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def locate_table(table, row_pos, col_pos, st):
	if len(table.rows) > row_pos:
		row = table.rows[row_pos]
		if len(row.cells) > col_pos:
			cell = row.cells[col_pos]
			for paragraph in cell.paragraphs:
				if st in paragraph.text:
					return True
	
	return False

def write_to_termsheet(cur, header, template, outfile):
	document = Document(template)

	for paragraph in document.paragraphs:
		if '[TODAY_DATE]' in paragraph.text:
			print paragraph.text
			paragraph.text = datetime.date.today().strftime("%d. %B %Y")

	for table in document.tables:
		if locate_table(table, 1, 0, "Liquidity Provider broker ID"):

			total_width = table.columns[0].width + table.columns[1].width
			tbl_style = table.style

			print table.style


			cur.execute("select * from termsheet")
			term_rows = cur.fetchall()
			term_row_cnt = len(term_rows)
			avg_width = total_width / term_row_cnt

			for i in range(1, term_row_cnt):
				table.add_column(avg_width)
				for row in table.rows:
					row.cells[i+1].add_paragraph(row.cells[i].text, row.cells[1].paragraphs[0].style)



		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					if '[Stock Code]' in paragraph.text:
						print paragraph.text
						paragraph.text = "26641"

	document.save(outfile)
	return

def merge_xlsx_db(cur, term_sheet_file, comp_info_file):
	header_term, arr_term = fl_util.xlsx_to_arr(term_sheet_file)
	header_info, arr_info = fl_util.xlsx_to_arr(comp_info_file)

	fl_util.create_tbl(cur, "term", header_term, arr_term)
	fl_util.create_tbl(cur, "info", header_info, arr_info)

	header_termsheet = fl_util.remove_dup(header_term, header_info)
	header_termsheet = '"'+'","'.join(header_term.split(','))+'"'

	cur.execute("create table termsheet (" + header_termsheet + ")")
	cur.execute("insert into termsheet ("+header_termsheet+") SELECT "+header_termsheet+" from term left join info on term.guarantor = info.equity where issuer <> '' ")

	cur.execute("drop table term")
	cur.execute("drop table info")

	info_rows = cur.execute("select * from termsheet")
	return header_termsheet

comp_info_file = "D:\\Projects\\HTI\\docx_automation\\Company Information.xlsx"
term_sheet_file = "D:\\Projects\\HTI\\docx_automation\\ClearedTermSheet_20160308_095932.csv.xlsx"
sld_eng_template = "D:\\Projects\\HTI\\docx_automation\\SLD (26774, 26775, 26778, 26779, 26780 english).docx"
sld_eng_output = "D:\\Projects\\HTI\\docx_automation\\SLD (English).docx"

conn, cur = fl_util.db_cur()
header_termsheet = merge_xlsx_db(cur, term_sheet_file, comp_info_file)

write_to_termsheet(cur, header_termsheet, sld_eng_template, sld_eng_output)